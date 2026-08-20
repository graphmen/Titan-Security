#!/usr/bin/env python3
"""Generate Titan Protection Phase 2 Enhancement & Security Monitoring Plan (.docx)."""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "_phase2_assets"
OUTPUT = DOCS / "Titan-Protection-Phase-2-Implementation-Plan.docx"

WORDMARK = ROOT / "web" / "public" / "emblem-wordmark.png"
APP_ICON = ROOT / "web" / "public" / "icons" / "icon-512.png"

BRAND = {
    "primary": RGBColor(0x1B, 0x43, 0x32),
    "success": RGBColor(0x40, 0x91, 0x6C),
    "text": RGBColor(0x0F, 0x1F, 0x17),
    "muted": RGBColor(0x3D, 0x5A, 0x48),
    "dimmed": RGBColor(0x7A, 0x94, 0x85),
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "primary_hex": "1B4332",
    "primary_light_hex": "D8F3DC",
    "success_hex": "40916C",
    "warning_hex": "F59E0B",
    "danger_hex": "EF4444",
    "surface_hex": "F4FAF6",
    "accent_hex": "2D6A4F",
}

# Matplotlib brand colours
MPL = {
    "primary": "#1B4332",
    "success": "#40916C",
    "light": "#D8F3DC",
    "surface": "#F4FAF6",
    "warning": "#F59E0B",
    "danger": "#EF4444",
    "muted": "#3D5A48",
    "accent": "#2D6A4F",
}


def hex_fill(cell, fill_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BRAND["text"]
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = BRAND["primary"]
        h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(6)


def setup_header_footer(doc: Document):
    logo = WORDMARK if WORDMARK.exists() else APP_ICON
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        hdr = section.header
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if logo.exists():
            hp.add_run().add_picture(str(logo), width=Inches(1.35))
        tr = hp.add_run("   Phase 2 — Enhancement & Security Monitoring Plan")
        tr.font.size = Pt(9)
        tr.font.color.rgb = BRAND["muted"]
        tr.italic = True
        ftr = section.footer
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(
            f"Titan Protection Operations Hub  ·  titanprotection.org  ·  "
            f"Arch Luviah Technologies © {date.today().year}  ·  Confidential"
        )
        r.font.size = Pt(8)
        r.font.color.rgb = BRAND["dimmed"]


def add_cover_page(doc: Document):
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.rows[0].cells[0]
    hex_fill(cell, BRAND["primary_hex"])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TITAN PROTECTION")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND["white"]

    doc.add_paragraph()
    if WORDMARK.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(str(WORDMARK), width=Inches(3.0))

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Phase 2 Implementation Plan")
    r.bold = True
    r.font.size = Pt(28)
    set_run_color(r, BRAND["primary"])

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Guard Monitoring Enhancements\n"
        "Security Operations · Patrol Verification · Real-Time Response"
    )
    r.font.size = Pt(13)
    set_run_color(r, BRAND["muted"])

    doc.add_paragraph()
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Document", "Phase 2 Enhancement & Security Monitoring Plan"),
        ("Version", "1.0"),
        ("Date", date.today().strftime("%d %B %Y")),
        ("Prepared for", "Titan Protection Security"),
        ("Live system", "https://titanprotection.org"),
    ]
    for i, (label, val) in enumerate(rows):
        hex_fill(meta.rows[i].cells[0], BRAND["primary_light_hex"])
        c0 = meta.rows[i].cells[0].paragraphs[0]
        c0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = c0.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        set_run_color(lr, BRAND["primary"])
        vr = meta.rows[i].cells[1].paragraphs[0].add_run(val)
        vr.font.size = Pt(10)

    doc.add_paragraph()
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("Built to Protect")
    r.italic = True
    r.font.size = Pt(12)
    set_run_color(r, BRAND["success"])
    doc.add_page_break()


def add_branded_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hex_fill(hdr[i], BRAND["primary_hex"])
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = BRAND["white"]
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            if ri % 2 == 1:
                hex_fill(cell, BRAND["surface_hex"])
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc, title, body, kind="note"):
    fill = BRAND["primary_light_hex"] if kind == "note" else "FEE2E2"
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    hex_fill(cell, fill)
    p = cell.paragraphs[0]
    tr = p.add_run(f"{title}  ")
    tr.bold = True
    set_run_color(tr, BRAND["primary"] if kind == "note" else RGBColor(0xB9, 0x1C, 0x1C))
    p.add_run(body)
    doc.add_paragraph()


def add_figure(doc, path: Path, caption: str, width=Inches(6.2)):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    set_run_color(r, BRAND["muted"])
    doc.add_paragraph()


def _save_fig(fig, name: str) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / name
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _box(ax, x, y, w, h, text, fc=MPL["light"], ec=MPL["primary"], fs=9, bold=False):
    patch = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.03,rounding_size=0.08",
        linewidth=1.5, edgecolor=ec, facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fs, color=MPL["primary"], fontweight="bold" if bold else "normal",
            wrap=True)


def _arrow(ax, x1, y1, x2, y2):
    ax.add_patch(FancyArrowPatch(
        (x1, y1), (x2, y2),
        arrowstyle="-|>", mutation_scale=12,
        linewidth=1.4, color=MPL["accent"],
    ))


def diagram_roadmap() -> Path:
    fig, ax = plt.subplots(figsize=(10, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5.5)
    ax.axis("off")
    ax.set_title("Titan Protection — Enhancement Roadmap", fontsize=14, fontweight="bold", color=MPL["primary"], pad=12)

    phases = [
        (0.3, 3.6, "PHASE 1\nFoundation", "Overdue patrol alerts\nGeofence exit (live)\nPush notifications\nOffline SOS queue", MPL["success"]),
        (2.7, 3.6, "PHASE 2\nCore Monitoring", "Real NFC patrol tags\nWelfare / dead-man checks\nGuard status board\nIncident escalation\nShift handover", MPL["primary"]),
        (5.1, 3.6, "PHASE 3\nTrust & Reporting", "Client portal\nPDF/CSV reports\nGPS anti-spoof\nProduction RLS", MPL["accent"]),
        (7.5, 3.6, "PHASE 4\nAdvanced", "CCTV / alarm hooks\nWhatsApp auto-alerts\nWebSocket live map\nAI anomaly detection", MPL["muted"]),
    ]
    for x, y, title, body, color in phases:
        _box(ax, x, y, 2.1, 0.65, title, fc=color, ec=color, fs=8, bold=True)
        for i, line in enumerate(body.split("\n")):
            ax.text(x + 1.05, y - 0.25 - i * 0.28, f"• {line}", ha="center", va="top", fontsize=7.5, color=MPL["muted"])

    for x in [2.5, 4.9, 7.3]:
        _arrow(ax, x, 3.95, x + 0.15, 3.95)

    _box(ax, 1.0, 0.5, 8.0, 0.9,
         "Goal: Prove guards are on site, patrolling on schedule, and responding to incidents — with audit-ready evidence.",
         fc=MPL["surface"], ec=MPL["success"], fs=9, bold=True)
    return _save_fig(fig, "01-roadmap.png")


def diagram_architecture() -> Path:
    fig, ax = plt.subplots(figsize=(10, 6.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    ax.set_title("Phase 2 — System Architecture Additions", fontsize=14, fontweight="bold", color=MPL["primary"], pad=12)

    _box(ax, 0.4, 5.0, 2.2, 0.8, "Titan Monitor\n(Guard Mobile)", fc="#E8F5E9")
    _box(ax, 3.5, 5.0, 2.2, 0.8, "Titan Supervisor\n(Mobile)", fc="#E8F5E9")
    _box(ax, 6.6, 5.0, 2.8, 0.8, "Command Centre\n(Web Admin)", fc="#E8F5E9")

    _box(ax, 2.5, 3.3, 5.0, 1.0, "Next.js API  /api/state  +  Supabase Postgres", fc=MPL["light"], bold=True)

    _box(ax, 0.3, 1.5, 2.0, 0.9, "Firebase FCM\n(Push — NEW)", fc="#FFF3E0", ec=MPL["warning"])
    _box(ax, 2.5, 1.5, 2.0, 0.9, "NFC Tag Service\n(NEW)", fc="#FFF3E0", ec=MPL["warning"])
    _box(ax, 4.7, 1.5, 2.0, 0.9, "Alert Engine\n(Enhanced)", fc="#FFF3E0", ec=MPL["warning"])
    _box(ax, 6.9, 1.5, 2.0, 0.9, "Audit & Reports\n(NEW)", fc="#FFF3E0", ec=MPL["warning"])

    _box(ax, 2.0, 0.2, 6.0, 0.75, "Supabase: guard_alerts · guard_attendance · checkpoints · occurrence_book · audit_log (new)", fc=MPL["surface"])

    for x in [1.5, 4.6, 8.0]:
        _arrow(ax, x, 5.0, x, 4.35)
    for x in [1.3, 3.5, 5.7, 7.9]:
        _arrow(ax, x, 3.3, x, 2.45)
    for x in [1.3, 3.5, 5.7, 7.9]:
        _arrow(ax, x, 1.5, x, 1.0)

    ax.text(5.0, 2.85, "Polling (10s) + Push (instant) for critical alerts", ha="center", fontsize=8, color=MPL["muted"], style="italic")
    return _save_fig(fig, "02-architecture.png")


def diagram_nfc_flow() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 4.5)
    ax.axis("off")
    ax.set_title("Real NFC Patrol Verification Flow", fontsize=13, fontweight="bold", color=MPL["primary"], pad=10)

    steps = [
        (0.3, 2.0, "1. Admin registers\nNFC tag ID\non patrol place"),
        (2.3, 2.0, "2. Guard taps phone\non physical tag\nat checkpoint"),
        (4.3, 2.0, "3. App reads tag UID\n+ captures GPS"),
        (6.3, 2.0, "4. Server validates\ntag + location\n+ schedule"),
        (8.0, 2.0, "5. OB log + score\nupdate + map\nrefresh"),
    ]
    for i, (x, y, t) in enumerate(steps):
        fc = MPL["light"] if i < 4 else "#C8E6C9"
        _box(ax, x, y, 1.7, 1.1, t, fc=fc, fs=8)
        if i < 4:
            _arrow(ax, x + 1.75, y + 0.55, x + 2.05, y + 0.55)

    _box(ax, 1.0, 0.4, 7.0, 0.75,
         "Security benefit: Physical proof of presence — cannot be faked with GPS spoofing alone",
         fc=MPL["surface"], ec=MPL["success"], fs=8.5, bold=True)
    return _save_fig(fig, "03-nfc-flow.png")


def diagram_push_notifications() -> Path:
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Push Notification Architecture (FCM)", fontsize=13, fontweight="bold", color=MPL["primary"], pad=10)

    _box(ax, 0.4, 3.5, 2.0, 0.9, "Alert Engine\n(missed clock-in,\nSOS, geofence)", fc=MPL["light"])
    _box(ax, 3.0, 3.5, 2.4, 0.9, "Notification\nDispatcher (NEW)", fc="#FFF3E0", ec=MPL["warning"], bold=True)
    _box(ax, 6.2, 3.5, 2.4, 0.9, "Firebase Cloud\nMessaging", fc="#E3F2FD", ec="#1565C0")

    _box(ax, 0.6, 1.5, 2.2, 0.8, "Supervisor\nMobile App", fc="#E8F5E9")
    _box(ax, 3.3, 1.5, 2.2, 0.8, "Guard\nMobile App", fc="#E8F5E9")
    _box(ax, 6.0, 1.5, 2.4, 0.8, "Web Command\nCentre", fc="#E8F5E9")

    _arrow(ax, 2.4, 3.95, 3.0, 3.95)
    _arrow(ax, 5.4, 3.95, 6.2, 3.95)
    _arrow(ax, 7.4, 3.5, 1.7, 2.35)
    _arrow(ax, 7.4, 3.5, 4.4, 2.35)
    _arrow(ax, 7.4, 3.5, 7.2, 2.35)

    ax.text(4.5, 0.5, "Target: SOS and critical alerts delivered in under 5 seconds (vs. 10s polling today)",
            ha="center", fontsize=8.5, color=MPL["muted"], style="italic")
    return _save_fig(fig, "04-push-notifications.png")


def diagram_geofence() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 4.8)
    ax.axis("off")
    ax.set_title("Geofence Exit Monitoring Flow", fontsize=13, fontweight="bold", color=MPL["primary"], pad=10)

    _box(ax, 0.4, 3.2, 2.2, 0.85, "Guard on duty\nGPS heartbeat\n(every 45s)", fc=MPL["light"])
    _box(ax, 3.0, 3.2, 2.4, 0.85, "Outside geofence\nfor > grace period?", fc="#FFF3E0", ec=MPL["warning"])
    _box(ax, 5.8, 3.5, 2.5, 0.7, "YES → Create alert\n+ push supervisor", fc="#FFCDD2", ec=MPL["danger"])
    _box(ax, 5.8, 2.3, 2.5, 0.7, "NO → Continue\nmonitoring", fc="#C8E6C9", ec=MPL["success"])

    _box(ax, 0.4, 1.0, 2.2, 0.85, "Guard re-enters\ngeofence", fc=MPL["light"])
    _box(ax, 3.0, 1.0, 2.4, 0.85, "Auto-resolve\nalert", fc="#C8E6C9", ec=MPL["success"])
    _box(ax, 5.8, 1.0, 2.5, 0.85, "Escalate if outside\n> 15 min (optional)", fc="#FFCDD2", ec=MPL["danger"])

    _arrow(ax, 2.6, 3.6, 3.0, 3.6)
    _arrow(ax, 5.4, 3.75, 5.8, 3.75)
    _arrow(ax, 5.4, 3.25, 5.8, 2.65)
    _arrow(ax, 2.6, 1.45, 3.0, 1.45)
    _arrow(ax, 5.4, 1.45, 5.8, 1.45)

    ax.text(4.5, 0.3, "Note: Geofence exit code exists today but is parked — Phase 1 enables with jitter protection",
            ha="center", fontsize=8, color=MPL["muted"], style="italic")
    return _save_fig(fig, "05-geofence.png")


def diagram_overdue_patrol() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 4.5)
    ax.axis("off")
    ax.set_title("Overdue Checkpoint / Patrol Alert Flow", fontsize=13, fontweight="bold", color=MPL["primary"], pad=10)

    _box(ax, 0.3, 2.5, 2.0, 1.0, "Checkpoint schedule\n(e.g. Every 2 hrs)\n+ last scan time", fc=MPL["light"])
    _box(ax, 2.7, 2.5, 2.2, 1.0, "Server evaluates\non each poll", fc=MPL["light"])
    _box(ax, 5.3, 2.8, 2.2, 0.75, "Overdue → Alert\nsupervisor", fc="#FFCDD2", ec=MPL["danger"])
    _box(ax, 5.3, 1.5, 2.2, 0.75, "On time → OK", fc="#C8E6C9", ec=MPL["success"])
    _box(ax, 0.3, 0.5, 7.2, 0.75,
         "Repeat every 30 min until scanned · Patrol compliance % on dashboard · GIS map layer",
         fc=MPL["surface"], ec=MPL["success"], fs=8.5, bold=True)

    _arrow(ax, 2.3, 3.0, 2.7, 3.0)
    _arrow(ax, 4.9, 3.15, 5.3, 3.15)
    _arrow(ax, 4.9, 2.85, 5.3, 1.95)
    return _save_fig(fig, "06-overdue-patrol.png")


def diagram_welfare() -> Path:
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Welfare Check / Dead Man's Switch", fontsize=13, fontweight="bold", color=MPL["primary"], pad=10)

    _box(ax, 0.4, 3.6, 2.2, 0.85, "Night shift starts\n(welfare mode on)", fc=MPL["light"])
    _box(ax, 3.0, 3.6, 2.4, 0.85, "Every 60–90 min:\n'Confirm OK' prompt", fc="#FFF3E0", ec=MPL["warning"])
    _box(ax, 6.0, 3.9, 2.5, 0.65, "Guard taps OK", fc="#C8E6C9", ec=MPL["success"])
    _box(ax, 6.0, 2.7, 2.5, 0.65, "No response 5 min", fc="#FFCDD2", ec=MPL["danger"])

    _box(ax, 0.4, 1.5, 2.2, 0.85, "Supervisor alert\n+ push notification", fc="#FFCDD2", ec=MPL["danger"])
    _box(ax, 3.0, 1.5, 2.4, 0.85, "Escalate to admin\nif no ack 10 min", fc="#FFCDD2", ec=MPL["danger"])
    _box(ax, 6.0, 1.5, 2.5, 0.85, "Log all checks\nin audit trail", fc=MPL["light"])

    _arrow(ax, 2.6, 4.0, 3.0, 4.0)
    _arrow(ax, 5.4, 4.05, 6.0, 4.05)
    _arrow(ax, 5.4, 3.75, 6.0, 3.05)
    _arrow(ax, 7.25, 2.7, 1.5, 2.4)
    _arrow(ax, 2.6, 1.95, 3.0, 1.95)
    _arrow(ax, 5.4, 1.95, 6.0, 1.95)

    ax.text(4.5, 0.4, "Ideal for lone guards at remote sites — extends existing no-movement alert system",
            ha="center", fontsize=8, color=MPL["muted"], style="italic")
    return _save_fig(fig, "07-welfare.png")


def diagram_status_board() -> Path:
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Live Guard Status Board (Command Centre)", fontsize=13, fontweight="bold", color=MPL["primary"], pad=10)

    headers = ["Guard", "Site", "Last GPS", "Last Patrol", "Movement", "Status"]
    col_x = [0.2, 1.5, 3.0, 4.5, 6.0, 7.5]
    for i, h in enumerate(headers):
        _box(ax, col_x[i], 4.0, 1.2 if i > 0 else 1.1, 0.5, h, fc=MPL["primary"], ec=MPL["primary"], fs=7, bold=True)
        for run in ax.texts:
            if run.get_text() == h:
                run.set_color("white")

    rows = [
        ["J. Dube", "Sandton HQ", "2 min ago", "45 min ago", "Active", "GREEN"],
        ["P. Nkosi", "Rosebank Mall", "1 min ago", "OVERDUE", "Still", "AMBER"],
        ["T. Moyo", "Midrand Estate", "8 min ago", "20 min ago", "Active", "GREEN"],
        ["S. Khumalo", "Pretoria Gate", "3 min ago", "OK", "Outside fence", "RED"],
    ]
    colours = {"GREEN": "#C8E6C9", "AMBER": "#FFF3E0", "RED": "#FFCDD2"}
    for ri, row in enumerate(rows):
        y = 3.2 - ri * 0.75
        status = row[-1]
        for ci, val in enumerate(row):
            w = 1.2 if ci > 0 else 1.1
            fc = colours.get(val, MPL["surface"] if ri % 2 else "white")
            if ci < 5:
                fc = MPL["surface"] if ri % 2 else "white"
            _box(ax, col_x[ci], y, w, 0.55, val, fc=fc, fs=7.5,
                 ec=MPL["danger"] if val == "RED" else MPL["warning"] if val == "AMBER" else MPL["primary"])

    ax.text(5.0, 0.35,
            "Status rules: GREEN = all OK · AMBER = overdue patrol or no movement · RED = geofence exit or missed welfare check",
            ha="center", fontsize=8, color=MPL["muted"], style="italic")
    return _save_fig(fig, "08-status-board.png")


def diagram_timeline() -> Path:
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.set_title("Suggested Implementation Timeline (Weeks)", fontsize=13, fontweight="bold", color=MPL["primary"], pad=10)

    blocks = [
        (0.5, 3.5, 3, "Phase 1\nWks 1–4", MPL["success"],
         ["Overdue patrol alerts", "Geofence exit live", "Push notifications", "Offline SOS"]),
        (4.0, 3.5, 4, "Phase 2\nWks 5–10", MPL["primary"],
         ["Real NFC scanning", "Welfare checks", "Status board", "Incident escalation", "Shift handover"]),
        (8.5, 3.5, 3.5, "Phase 3\nWks 11–16", MPL["accent"],
         ["Client portal", "PDF reports", "GPS anti-spoof", "Production RLS"]),
        (12.5, 3.5, 3, "Phase 4\nWks 17+", MPL["muted"],
         ["CCTV hooks", "WhatsApp auto", "WebSocket map"]),
    ]
    for x, y, w, title, color, items in blocks:
        rect = FancyBboxPatch((x, y), w, 0.7, boxstyle="round,pad=0.02", facecolor=color, edgecolor=color)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + 0.35, title, ha="center", va="center", fontsize=8, color="white", fontweight="bold")
        for i, item in enumerate(items):
            ax.text(x + 0.15, y - 0.25 - i * 0.28, f"• {item}", fontsize=7, color=MPL["muted"])

    ax.plot([0.5, 15.5], [2.0, 2.0], color=MPL["primary"], linewidth=2)
    for wk in range(1, 17):
        x = 0.5 + (wk - 1) * (15 / 16)
        ax.plot([x, x], [1.9, 2.1], color=MPL["primary"], linewidth=1)
        if wk % 2 == 1:
            ax.text(x, 1.65, str(wk), ha="center", fontsize=6, color=MPL["muted"])
    ax.text(8.0, 1.35, "Weeks →", ha="center", fontsize=8, color=MPL["muted"])
    return _save_fig(fig, "09-timeline.png")


def diagram_incident_escalation() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 4.8)
    ax.axis("off")
    ax.set_title("Incident Escalation Workflow", fontsize=13, fontweight="bold", color=MPL["primary"], pad=10)

    _box(ax, 0.4, 3.3, 1.8, 0.8, "Guard logs\nincident", fc=MPL["light"])
    _box(ax, 2.6, 3.3, 2.0, 0.8, "Auto-assign\nterritory supervisor", fc="#FFF3E0", ec=MPL["warning"])
    _box(ax, 5.0, 3.5, 1.8, 0.6, "Low severity\nInvestigating", fc="#C8E6C9")
    _box(ax, 5.0, 2.5, 1.8, 0.6, "High severity\nImmediate push", fc="#FFCDD2", ec=MPL["danger"])
    _box(ax, 7.2, 3.3, 1.5, 0.8, "SLA timer\n15 min", fc=MPL["light"])
    _box(ax, 2.6, 1.2, 2.0, 0.8, "Unassigned?\nAlert admin", fc="#FFCDD2", ec=MPL["danger"])
    _box(ax, 5.0, 1.2, 2.0, 0.8, "Resolved +\nOB archive", fc="#C8E6C9", ec=MPL["success"])

    _arrow(ax, 2.2, 3.7, 2.6, 3.7)
    _arrow(ax, 4.6, 3.85, 5.0, 3.75)
    _arrow(ax, 4.6, 3.45, 5.0, 2.85)
    _arrow(ax, 6.8, 3.7, 7.2, 3.7)
    _arrow(ax, 8.0, 3.3, 3.6, 2.05)
    _arrow(ax, 4.6, 1.6, 5.0, 1.6)
    return _save_fig(fig, "10-incident-escalation.png")


def build_document():
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "roadmap": diagram_roadmap(),
        "architecture": diagram_architecture(),
        "nfc": diagram_nfc_flow(),
        "push": diagram_push_notifications(),
        "geofence": diagram_geofence(),
        "overdue": diagram_overdue_patrol(),
        "welfare": diagram_welfare(),
        "status_board": diagram_status_board(),
        "timeline": diagram_timeline(),
        "incident": diagram_incident_escalation(),
    }

    doc = Document()
    setup_styles(doc)
    setup_header_footer(doc)
    add_cover_page(doc)

    # ── TOC ──
    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Executive Summary",
        "2. Current Platform Baseline",
        "3. Strategic Objectives",
        "4. Enhancement Roadmap Overview",
        "5. Phase 1 — Foundation (Prerequisites)",
        "6. Phase 2 — Core Guard Monitoring (Detailed)",
        "7. Phase 3 — Trust, Reporting & Compliance",
        "8. Phase 4 — Advanced Integrations",
        "9. System Architecture",
        "10. Feature Specifications & Process Flows",
        "11. Implementation Timeline",
        "12. Success Metrics & KPIs",
        "13. Risks & Mitigations",
        "14. Appendix — Priority Matrix",
    ]
    add_numbered(doc, toc)
    doc.add_page_break()

    # ── 1 Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document defines the Phase 2 enhancement plan for the Titan Protection Operations Hub — "
        "a integrated platform comprising the web Command Centre, Supervisor Portal, Titan Monitor guard "
        "mobile app, and Titan Supervisor mobile app. The plan focuses on strengthening guard monitoring, "
        "increasing site security, and providing audit-ready evidence for clients and PSIRA compliance."
    )
    add_callout(doc, "Purpose.",
                "Transform Titan Protection from a operational tracking platform into a proactive security "
                "monitoring system that proves guards are on site, patrolling on schedule, and responding to incidents.")
    doc.add_paragraph(
        "The enhancements are organised into four implementation phases. Phase 1 establishes real-time "
        "alerting foundations. Phase 2 (the primary focus of this document) delivers core monitoring capabilities "
        "including physical NFC patrol verification, welfare checks, live guard status boards, and structured "
        "incident escalation. Phases 3 and 4 extend the platform with client reporting, anti-fraud measures, "
        "and advanced integrations."
    )
    doc.add_page_break()

    # ── 2 Baseline ──
    doc.add_heading("2. Current Platform Baseline", level=1)
    doc.add_paragraph("The following capabilities are already live in production at titanprotection.org:")
    add_branded_table(doc, ["Module", "Current capability"], [
        ["Guard mobile (Titan Monitor)", "Clock in/out with geofence · Patrol tap · SOS · Incidents · Checklists · Visitors · Offline queue"],
        ["Supervisor mobile", "Territory-scoped guards · Shift roster · Alert dismiss · SOS clear · Premises registration"],
        ["Web Command Centre", "GIS map · Live attendance · Occurrence Book · Access desk · Checklist builder · Siren alerts"],
        ["Alerts (live)", "Missed clock-in/out (30 min repeat) · No movement · License expiry · Shift swap"],
        ["Alerts (parked)", "Geofence exit (code exists, disabled in settings)"],
        ["Integrations", "Email PIN delivery (Resend) · WhatsApp manual links · HTTP polling (~10s)"],
        ["Not yet built", "Push notifications · Real NFC hardware · Overdue patrol enforcement · Client portal · PDF reports"],
    ])

    # ── 3 Objectives ──
    doc.add_heading("3. Strategic Objectives", level=1)
    add_bullets(doc, [
        "Prove guard presence — physical and GPS-verified, not honour-system check-ins.",
        "Enforce patrol schedules — alert when checkpoints are overdue, not just display progress.",
        "Reduce incident response time — instant push for SOS, geofence exit, and missed clock-in.",
        "Protect lone guards — welfare checks and dead-man's switch on night shifts.",
        "Win and retain clients — compliance dashboards and monthly audit reports.",
        "Meet PSIRA expectations — license tracking, training records, and defensible occurrence logs.",
    ])

    # ── 4 Roadmap ──
    doc.add_heading("4. Enhancement Roadmap Overview", level=1)
    add_figure(doc, diagrams["roadmap"],
               "Figure 1 — Four-phase enhancement roadmap from foundation alerts to advanced integrations")
    doc.add_page_break()

    # ── 5 Phase 1 ──
    doc.add_heading("5. Phase 1 — Foundation (Prerequisites)", level=1)
    doc.add_paragraph(
        "Phase 1 items are prerequisites that should be completed before or in parallel with Phase 2. "
        "They address the highest-frequency monitoring gaps using existing code where possible."
    )
    add_branded_table(doc, ["Feature", "Description", "Effort", "Security impact"], [
        ["Overdue patrol alerts", "Server alert when checkpoint scan exceeds schedule (e.g. Every 2 hrs) + 30 min repeat", "Medium", "High — proves patrols happened"],
        ["Geofence exit (live)", "Re-enable with grace period to avoid GPS jitter; auto-resolve on re-entry", "Low", "High — proves guard stayed on site"],
        ["Push notifications (FCM)", "Firebase Cloud Messaging for SOS, critical alerts, missed clock-in", "Medium", "High — sub-5s response vs 10s poll"],
        ["Offline SOS queue", "Queue panic alert locally; auto-send when connectivity returns", "Low", "Critical — guard safety in dead zones"],
    ])
    add_figure(doc, diagrams["geofence"], "Figure 2 — Geofence exit monitoring with grace period and auto-resolve")
    add_figure(doc, diagrams["overdue"], "Figure 3 — Overdue checkpoint evaluation and recurring supervisor alerts")
    add_figure(doc, diagrams["push"], "Figure 4 — Push notification architecture replacing poll-only delivery")

    # ── 6 Phase 2 ──
    doc.add_heading("6. Phase 2 — Core Guard Monitoring (Detailed)", level=1)
    doc.add_paragraph(
        "Phase 2 is the primary focus of this implementation plan. These features transform patrol tracking "
        "into verifiable security monitoring and give supervisors real-time situational awareness."
    )

    doc.add_heading("6.1 Real NFC Patrol Tag Scanning", level=2)
    doc.add_paragraph(
        "Today, patrol checkpoints are confirmed via an in-app button tap. Phase 2 integrates physical NFC tags "
        "installed at doors, gates, and patrol points. The guard's phone reads the tag's unique ID, which the server "
        "validates against the registered nfc_code on each patrol place."
    )
    add_bullets(doc, [
        "Admin registers NFC tag UID when creating or editing a patrol place.",
        "Guard taps phone on tag — Capacitor NFC plugin reads UID + captures GPS.",
        "Server validates: tag ID matches · guard is on duty · location within geofence · schedule met.",
        "Failed validation logged to Occurrence Book with reason (wrong tag, wrong site, etc.).",
    ])
    add_figure(doc, diagrams["nfc"], "Figure 5 — Real NFC patrol verification end-to-end flow")

    doc.add_heading("6.2 Welfare Check / Dead Man's Switch", level=2)
    doc.add_paragraph(
        "Extends the existing no-movement alert into a formal welfare protocol for lone guards on night shifts "
        "or remote sites. The guard app prompts 'Confirm OK' every 60–90 minutes (configurable per site)."
    )
    add_bullets(doc, [
        "No response within 5 minutes → supervisor alert + push notification.",
        "No supervisor acknowledgement within 10 minutes → escalate to Master Admin.",
        "All welfare checks logged in audit trail with timestamps.",
        "Can be enabled per-premise (high-risk sites only) or per shift type (night only).",
    ])
    add_figure(doc, diagrams["welfare"], "Figure 6 — Welfare check and escalation flow")

    doc.add_heading("6.3 Live Guard Status Board", level=2)
    doc.add_paragraph(
        "A dedicated Command Centre view showing every on-duty guard as a live status card. "
        "Supervisors get at-a-glance situational awareness without opening the GIS map."
    )
    add_branded_table(doc, ["Status", "Colour", "Trigger conditions"], [
        ["GREEN", "Normal", "On duty · GPS fresh · patrol on schedule · inside geofence"],
        ["AMBER", "Warning", "Overdue patrol OR no movement beyond threshold"],
        ["RED", "Critical", "Geofence exit OR missed welfare check OR active SOS"],
    ])
    add_figure(doc, diagrams["status_board"], "Figure 7 — Live Guard Status Board mock-up for Command Centre")

    doc.add_heading("6.4 Shift Handover Verification", level=2)
    doc.add_paragraph(
        "Prevents unstaffed gaps between shifts at client sites. Outgoing guard must clock out with handover notes; "
        "incoming guard must acknowledge before patrol begins."
    )
    add_bullets(doc, [
        "Outgoing guard: clock out + mandatory handover notes (relief guard name, site status, incidents).",
        "Incoming guard: acknowledge handover within 15 minutes of shift start.",
        "Gap alert: if no guard on duty for > X minutes between shifts → supervisor notification.",
        "Handover record stored in Occurrence Book and visible on Guard Status Board.",
    ])

    doc.add_heading("6.5 Site-Specific Alert Rules", level=2)
    doc.add_paragraph("Not every client site has the same risk profile. Phase 2 adds per-premise configuration overrides:")
    add_branded_table(doc, ["Setting", "Example use"], [
        ["Geofence radius override", "Large estate: 8m · Small office: 5m"],
        ["Patrol frequency", "High-risk warehouse: every 1 hr · Office park: every 2 hrs"],
        ["Welfare check interval", "Remote gate house: every 60 min · Mall patrol: disabled"],
        ["Night shift strict mode", "Stricter no-movement + mandatory welfare on 18:00–06:00 shifts"],
    ])

    doc.add_heading("6.6 Incident Escalation Workflow", level=2)
    doc.add_paragraph(
        "Incidents are logged today but lack structured escalation. Phase 2 adds auto-assignment, SLA timers, "
        "and severity-based routing."
    )
    add_bullets(doc, [
        "Auto-assign to territory supervisor on incident creation.",
        "Low severity: status workflow Unassigned → Investigating → Resolved.",
        "High severity (assault, fire, theft): immediate push to supervisor + admin.",
        "SLA alert if unassigned > 15 minutes.",
        "Full photo and voice memo playback on web Occurrence Book.",
    ])
    add_figure(doc, diagrams["incident"], "Figure 8 — Incident escalation and SLA workflow")
    doc.add_page_break()

    # ── 7 Phase 3 ──
    doc.add_heading("7. Phase 3 — Trust, Reporting & Compliance", level=1)
    add_branded_table(doc, ["Feature", "Description", "Business value"], [
        ["Client portal (read-only)", "Live guards on site · patrol completion · incident summary", "Sales differentiator"],
        ["Automated PDF/CSV reports", "Weekly/monthly compliance per client site", "Contract retention"],
        ["GPS anti-spoof detection", "Reject mock locations · flag teleport events", "Fraud prevention"],
        ["Production RLS hardening", "Tenant isolation · role-based Supabase policies", "Go-live requirement"],
        ["Audit trail module", "Who dismissed alerts · export OB + attendance", "Legal defensibility"],
        ["PSIRA compliance dashboard", "License + training expiry · site staffing blocks", "Regulatory compliance"],
    ])

    # ── 8 Phase 4 ──
    doc.add_heading("8. Phase 4 — Advanced Integrations", level=1)
    add_branded_table(doc, ["Feature", "Description"], [
        ["CCTV / alarm integration", "OB entry when site alarm triggers · map camera to premise"],
        ["WhatsApp auto-alerts to clients", "'Guard clocked in at your site' using existing WhatsApp infra"],
        ["WebSocket live map", "Sub-second guard positions on GIS for control-room displays"],
        ["AI patrol anomaly detection", "Unusual movement patterns · repeatedly skipped zones"],
        ["Body-worn camera upload", "Incident evidence chain"],
        ["Multi-language mobile", "English · Afrikaans · Zulu for guard UX"],
    ])

    # ── 9 Architecture ──
    doc.add_heading("9. System Architecture", level=1)
    doc.add_paragraph(
        "Phase 2 additions integrate with the existing Next.js / Supabase stack. No platform replacement is required. "
        "New components (FCM dispatcher, NFC validation, enhanced alert engine, audit log) extend the current /api/state "
        "action model and relational database schema."
    )
    add_figure(doc, diagrams["architecture"],
               "Figure 9 — Phase 2 architecture showing new services alongside existing platform components")
    doc.add_page_break()

    # ── 10 Feature specs ──
    doc.add_heading("10. Feature Specifications & Process Flows", level=1)
    doc.add_paragraph(
        "The following table summarises each Phase 2 feature with its primary user, trigger, and expected outcome."
    )
    add_branded_table(doc, ["Feature", "Primary user", "Trigger", "Outcome"], [
        ["Real NFC patrol", "Guard", "Phone tap on physical tag", "Verified checkpoint scan in OB + map"],
        ["Welfare check", "Guard / Supervisor", "Timer (60–90 min)", "OK logged or escalation alert"],
        ["Status board", "Admin / Supervisor", "Continuous poll + push", "Green/Amber/Red guard cards"],
        ["Shift handover", "Guard / Supervisor", "Shift change window", "No unstaffed gap at site"],
        ["Site alert rules", "Admin", "Per-premise config", "Tailored monitoring per client site"],
        ["Incident escalation", "Guard / Supervisor", "Incident logged", "Auto-assign + SLA + push"],
    ])

    # ── 11 Timeline ──
    doc.add_heading("11. Implementation Timeline", level=1)
    doc.add_paragraph(
        "The suggested timeline assumes one development team working on the existing Titan Protection codebase. "
        "Phases can overlap where dependencies allow (e.g. push notifications in Phase 1 unblock Phase 2 welfare alerts)."
    )
    add_figure(doc, diagrams["timeline"], "Figure 10 — Suggested 16+ week implementation timeline across all phases")
    add_branded_table(doc, ["Phase", "Duration", "Deliverables"], [
        ["Phase 1 — Foundation", "Weeks 1–4", "Overdue patrol · Geofence live · FCM push · Offline SOS"],
        ["Phase 2 — Core Monitoring", "Weeks 5–10", "NFC · Welfare · Status board · Handover · Escalation"],
        ["Phase 3 — Trust & Reporting", "Weeks 11–16", "Client portal · Reports · Anti-spoof · RLS · Audit"],
        ["Phase 4 — Advanced", "Week 17+", "CCTV · WhatsApp auto · WebSocket · AI anomaly"],
    ])

    # ── 12 KPIs ──
    doc.add_heading("12. Success Metrics & KPIs", level=1)
    add_branded_table(doc, ["Metric", "Target", "Measurement"], [
        ["SOS notification delivery", "< 5 seconds", "FCM delivery timestamp vs trigger"],
        ["Patrol compliance rate", "> 95% per site", "Scans on time / scheduled scans"],
        ["False geofence alerts", "< 2% of shifts", "Dismissed geofence / total geofence alerts"],
        ["Missed clock-in detection", "100% of scheduled shifts", "Alerts fired / shifts started"],
        ["Incident assignment time", "< 5 minutes average", "Created → Investigating timestamp"],
        ["Client report delivery", "Automated monthly", "PDF generated and emailed on schedule"],
        ["Guard welfare response rate", "> 98%", "OK taps / welfare prompts"],
    ])

    # ── 13 Risks ──
    doc.add_heading("13. Risks & Mitigations", level=1)
    add_branded_table(doc, ["Risk", "Impact", "Mitigation"], [
        ["GPS inaccuracy indoors", "False geofence / patrol alerts", "NFC for indoor checkpoints · grace periods · accuracy thresholds"],
        ["Guard phone battery dies", "Monitoring blind spot", "Low-battery alert · supervisor notification · status board RED"],
        ["NFC tag tampering", "Fake patrol scans", "Tag UID registration · GPS co-validation · audit log"],
        ["Push notification opt-out", "Missed critical alerts", "Fallback to polling + SMS for SOS · mandatory FCM for supervisors"],
        ["Scope creep across phases", "Delayed go-live", "Strict phase gates · Phase 2 complete before Phase 3 starts"],
        ["Supabase RLS in production", "Data leak between tenants", "Phase 3 RLS hardening before multi-client rollout"],
    ])

    # ── 14 Appendix ──
    doc.add_heading("14. Appendix — Priority Matrix", level=1)
    doc.add_paragraph(
        "Features ranked by security impact versus implementation effort. "
        "Start with top-right quadrant (high impact, lower effort)."
    )
    add_branded_table(doc, ["Priority", "Feature", "Impact", "Effort", "Phase"], [
        ["1", "Overdue patrol alerts", "High", "Medium", "1"],
        ["2", "Geofence exit (live)", "High", "Low", "1"],
        ["3", "Push notifications", "High", "Medium", "1"],
        ["4", "Real NFC tags", "High", "Medium", "2"],
        ["5", "Live status board", "High", "Medium", "2"],
        ["6", "Welfare checks", "High", "Medium", "2"],
        ["7", "Incident escalation", "Medium", "Medium", "2"],
        ["8", "Client portal + reports", "Medium", "High", "3"],
        ["9", "GPS anti-spoof", "Medium", "High", "3"],
        ["10", "CCTV integration", "Medium", "Very high", "4"],
    ])

    add_callout(doc, "Next step.",
                "Review and approve this plan, then begin Phase 1 foundation work while procuring NFC tags "
                "for Phase 2 pilot sites. Recommended pilot: one high-traffic site with 4–6 patrol points.")

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    build_document()
