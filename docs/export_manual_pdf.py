"""Export Titan Protection User Manual to PDF via fpdf2."""
from docx import Document
from fpdf import FPDF
from pathlib import Path

DOCX = Path(__file__).parent / "Titan-Protection-User-Manual.docx"
PDF_OUT = Path(__file__).parent / "Titan-Protection-User-Manual.pdf"

GREEN = (27, 67, 50)
DARK = (15, 31, 23)


def clean(text):
    return (
        text.replace("\u2014", " - ")
        .replace("\u2013", "-")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2022", "-")
        .encode("latin-1", "replace")
        .decode("latin-1")
    )


class ManualPDF(FPDF):
    def header(self):
        if self.page_no() > 1:
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(*GREEN)
            self.cell(0, 8, "Titan Protection - User Manual", align="L")
            self.ln(4)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 8, f"Page {self.page_no()} | Built to Protect", align="C")


def add_table(pdf, table):
    if not table.rows:
        return
    cols = len(table.rows[0].cells)
    width = (pdf.w - pdf.l_margin - pdf.r_margin) / cols
    for r_idx, row in enumerate(table.rows):
        if pdf.get_y() > pdf.h - 25:
            pdf.add_page()
        pdf.set_font("Helvetica", "B" if r_idx == 0 else "", 8)
        pdf.set_text_color(*(GREEN if r_idx == 0 else DARK))
        if r_idx == 0:
            pdf.set_fill_color(232, 245, 233)
        else:
            pdf.set_fill_color(255, 255, 255)
        for cell in row.cells:
            pdf.cell(width, 7, clean(cell.text.strip())[:80], border=1, fill=True)
        pdf.ln(7)


def build_pdf():
    doc = Document(DOCX)
    pdf = ManualPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]

        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(block, doc)
            text = para.text.strip()
            if not text:
                pdf.ln(3)
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                pdf.ln(6)
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_text_color(*GREEN)
            elif "Heading 2" in style:
                pdf.ln(4)
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(*GREEN)
            elif "Heading 3" in style:
                pdf.ln(3)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(*DARK)
            elif "List" in style:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(*DARK)
                text = "- " + text
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(*DARK)
            pdf.multi_cell(0, 5.5, clean(text))
            pdf.ln(1)

        elif tag == "tbl":
            from docx.table import Table
            pdf.ln(3)
            add_table(pdf, Table(block, doc))
            pdf.ln(4)

    pdf.output(str(PDF_OUT))
    print(f"Created: {PDF_OUT} ({PDF_OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build_pdf()
