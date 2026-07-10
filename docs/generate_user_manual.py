"""Generate Titan Protection User Manual (Word .docx)."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUTPUT = Path(__file__).parent / "Titan-Protection-User-Manual.docx"
GREEN = RGBColor(0x1B, 0x43, 0x32)
DARK = RGBColor(0x0F, 0x1F, 0x17)


def set_heading_style(doc):
    for i in range(1, 4):
        style = doc.styles[f"Heading {i}"]
        style.font.color.rgb = GREEN
        style.font.bold = True


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()
    set_heading_style(doc)

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TITAN PROTECTION\n")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = GREEN
    run2 = title.add_run("Security Operations Platform\nUser Manual")
    run2.font.size = Pt(18)
    run2.font.color.rgb = DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Built to Protect\n").italic = True
    sub.add_run("\nVersion 1.0  |  July 2026\n")
    sub.add_run("Developed by Arch Luviah Technologies")

    doc.add_page_break()

    # --- TOC ---
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. System Overview",
        "3. Getting Started",
        "4. Web Dashboard — Command Centre",
        "5. Web Dashboard — Master Administration",
        "6. Mobile Guard Application",
        "7. End-to-End Workflows",
        "8. Offline Mode & Synchronisation",
        "9. Multi-Tenant Management",
        "10. Troubleshooting & FAQ",
        "Appendix A: Video Walkthrough Script",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # --- 1 Introduction ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Titan Protection is a multi-tenant security operations platform designed for "
        "security companies, facility managers, and on-site guard teams. The system "
        "combines a web-based Command Centre for supervisors with a mobile Guard "
        "Application for field officers."
    )
    doc.add_paragraph(
        "Together, these applications enable real-time patrol monitoring, emergency SOS "
        "alerts, incident logging (Occurrence Book), compliance checklists, and visitor "
        "access control — all synchronised through a central backend."
    )

    # --- 2 Overview ---
    doc.add_heading("2. System Overview", level=1)
    doc.add_heading("2.1 Components", level=2)
    add_table(doc,
        ["Component", "Platform", "Primary Users"],
        [
            ["Command Centre (Web)", "Browser (Next.js)", "Supervisors, control room operators, administrators"],
            ["Guard Mobile App", "Browser / Android (Capacitor)", "Security guards, access desk officers"],
            ["Backend API", "Cloud / local server", "Connects web and mobile in real time"],
        ],
    )
    doc.add_heading("2.2 Key Capabilities", level=2)
    for item in [
        "Live patrol checkpoint tracking with NFC tap confirmation",
        "Emergency SOS panic button with instant Command Centre alerts",
        "Digital Occurrence Book (OB) for incidents and patrol logs",
        "Customisable safety and compliance checklists",
        "Visitor and vehicle access registration",
        "Multi-tenant support for multiple security providers or client sites",
        "Offline operation on mobile with automatic sync when connection restores",
    ]:
        add_bullet(doc, item)

    # --- 3 Getting Started ---
    doc.add_heading("3. Getting Started", level=1)
    doc.add_heading("3.1 Web Dashboard Access", level=2)
    add_numbered(doc, "Open a modern web browser (Chrome, Edge, or Firefox recommended).")
    add_numbered(doc, "Navigate to your Command Centre URL (e.g. http://localhost:3001 for development).")
    add_numbered(doc, "The dashboard loads automatically and displays live operational data.")
    doc.add_paragraph()

    doc.add_heading("3.2 Mobile App Access", level=2)
    add_numbered(doc, "Open the Guard Mobile app in a browser (e.g. http://localhost:5174) or launch the Android APK.")
    add_numbered(doc, "Ensure the device is connected to the same network as the Command Centre server.")
    add_numbered(doc, "Leave the Server URL empty for automatic proxy in development, or enter the full backend URL (e.g. http://192.168.1.10:3001) for production or physical devices.")
    add_numbered(doc, "Confirm the Wi-Fi icon shows green (online). Tap it to toggle offline mode if needed.")
    doc.add_paragraph()

    doc.add_heading("3.3 Recommended Setup", level=2)
    add_table(doc,
        ["Role", "Application", "Action"],
        [
            ["Supervisor", "Web Command Centre", "Monitor patrols, respond to SOS, manage visitors"],
            ["Guard", "Mobile App", "Perform patrol taps, report incidents, check in visitors"],
            ["Administrator", "Web Master Admin", "Create tenants, publish checklist templates"],
        ],
    )

    doc.add_page_break()

    # --- 4 Web Command Centre ---
    doc.add_heading("4. Web Dashboard — Command Centre", level=1)
    doc.add_paragraph(
        "The Command Centre is the primary operations view for supervisors. Access it by "
        "selecting Command Centre in the left sidebar."
    )

    doc.add_heading("4.1 Dashboard Layout", level=2)
    for item in [
        "Left sidebar — Titan Protection branding, navigation, tenant selector, and siren toggle",
        "Top header — page title, connection status badge (Demo Mode or Live Database)",
        "Statistics cards — Compliance Rate, Taps Scanned, Active Guests, Active Alerts",
        "Live Guard Geofence Tracker — animated map showing checkpoint and guard positions",
        "Occurrence Book (OB) — scrolling feed of all security events",
        "Access Desk — visitor check-in form and live visitor log table",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4.2 Statistics Cards", level=2)
    add_table(doc,
        ["Card", "Description"],
        [
            ["Compliance Rate", "Percentage of checkpoints scanned vs. total assigned checkpoints"],
            ["Taps Scanned", "Number of NFC checkpoint scans completed (e.g. 2 / 4)"],
            ["Active Guests", "Visitors currently checked in on site"],
            ["Active Alerts", "Open intrusion alerts or unassigned incidents"],
        ],
    )

    doc.add_heading("4.3 Live Guard Geofence Tracker", level=2)
    doc.add_paragraph(
        "The map canvas displays facility checkpoints as coloured nodes. Green nodes indicate "
        "scanned checkpoints; grey nodes are pending. A blue guard dot animates along the patrol "
        "path to simulate live GPS movement. Use Reset Patrol to clear all scan statuses for the active tenant."
    )

    doc.add_heading("4.4 Occurrence Book (OB)", level=2)
    doc.add_paragraph("Every security event appears in the OB feed, including:")
    for item in [
        "Patrol Tap — guard checked in at an NFC checkpoint",
        "SOS Panic Alarm — emergency distress signal (highlighted in red)",
        "Intrusion Alert — security incident reported by a guard",
        "Checklist Submission — completed compliance audit",
        "General Patrol — shift and routine activity logs",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph(
        "Each entry shows the event type, timestamp, guard name, description, and optional "
        "photo/voice attachments. Supervisors can change incident status using the dropdown: "
        "Unassigned → Investigating → Resolved."
    )

    doc.add_heading("4.5 SOS Panic Alerts", level=2)
    doc.add_paragraph(
        "When a guard triggers SOS on mobile, a red Critical Distress Signal banner appears "
        "at the top of the Command Centre. An audible siren plays (if Siren Alerts is enabled "
        "in the sidebar). Click Dismiss Alarm after the situation is handled."
    )

    doc.add_heading("4.6 Visitor Management", level=2)
    doc.add_paragraph("Use the Access Desk Registry form to manually check in guests:")
    for item in [
        "Visitor Full Name (required)",
        "National ID / Passport (required)",
        "Representing company (optional)",
        "Vehicle license plate (optional)",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph(
        "The Live Visitor & Vehicle Log table shows all check-ins. Click Check Out when a visitor leaves."
    )

    doc.add_heading("4.7 Tenant Switching", level=2)
    doc.add_paragraph(
        "Use the Active Security Tenant dropdown in the sidebar to switch between security "
        "providers (e.g. Titan Protection, Alpha Guard, Omega Watchmen). All dashboard data "
        "filters to the selected tenant."
    )

    doc.add_page_break()

    # --- 5 Master Admin ---
    doc.add_heading("5. Web Dashboard — Master Administration", level=1)
    doc.add_paragraph(
        "Select Master Admin in the sidebar to configure the platform."
    )

    doc.add_heading("5.1 Checklist Template Builder", level=2)
    doc.add_paragraph("Create custom safety audits for guards to complete on mobile:")
    add_numbered(doc, "Enter a Checklist Name and Guidelines/Instructions.")
    add_numbered(doc, "Add questions using the Add button. Each question can be Yes/No or Text Box.")
    add_numbered(doc, "Click Publish Checklist Template.")
    doc.add_paragraph("Published templates appear in the Active Checklists list and become available on guard devices.")

    doc.add_heading("5.2 Onboard Security Provider", level=2)
    doc.add_paragraph(
        "Register a new tenant (security company or client site) by entering the Agency Name, "
        "selecting a branding colour, and clicking Onboard. A default Reception Desk checkpoint "
        "is created automatically."
    )

    doc.add_heading("5.3 Completed Audits", level=2)
    doc.add_paragraph(
        "Review all checklist submissions from guards, including answers to each question, "
        "timestamp, and guard name."
    )

    doc.add_page_break()

    # --- 6 Mobile ---
    doc.add_heading("6. Mobile Guard Application", level=1)
    doc.add_paragraph(
        "The Guard Mobile app is optimised for field use with four main tabs at the bottom."
    )

    doc.add_heading("6.1 Header & Settings", level=2)
    for item in [
        "Titan Protection logo and Built to Protect tagline",
        "Wi-Fi icon — green when online; tap to switch offline mode",
        "Offline sync badge — shows queued items; tap to sync manually",
        "Active Guard name — tap to edit your officer name",
        "Distress SOS button — triggers emergency alert",
        "Server Network Config — set backend URL and tenant profile",
    ]:
        add_bullet(doc, item)

    doc.add_heading("6.2 Patrol Tab", level=2)
    doc.add_paragraph(
        "The radar map shows nearby checkpoints relative to the guard position. Below the map, "
        "each checkpoint displays its name, NFC code, schedule frequency, and scan status."
    )
    doc.add_paragraph("To record a patrol tap:")
    add_numbered(doc, "Approach the physical NFC checkpoint.")
    add_numbered(doc, "Tap the NFC Tap button next to the checkpoint name.")
    add_numbered(doc, "A confirmation toast appears and the button changes to Tap OK (green).")
    add_numbered(doc, "The Command Centre updates within seconds.")

    doc.add_heading("6.3 Incident OB Tab", level=2)
    doc.add_paragraph("File an Occurrence Book entry:")
    add_numbered(doc, "Select an Incident Category (Intrusion Alert, Safety Hazard, Theft Report, etc.).")
    add_numbered(doc, "Enter a detailed description.")
    add_numbered(doc, "Optionally attach a photo (Capture Image) or voice memo (Voice Memo).")
    add_numbered(doc, "Tap Submit Occurrence Log.")
    doc.add_paragraph("The incident appears immediately on the Command Centre OB feed.")

    doc.add_heading("6.4 Checklist Tab", level=2)
    add_numbered(doc, "Select a published checklist template and tap Inspect.")
    add_numbered(doc, "Answer each question (Yes/No buttons or text fields).")
    add_numbered(doc, "Track progress via the Audit Progress bar.")
    add_numbered(doc, "Tap Submit Completed Audit.")

    doc.add_heading("6.5 Access Tab", level=2)
    add_numbered(doc, "Tap Scan Visitor QR Code to simulate badge scanning (auto-fills visitor details).")
    add_numbered(doc, "Complete or verify the visitor form fields.")
    add_numbered(doc, "Tap Register Access Entry.")
    doc.add_paragraph("The visitor appears on the web dashboard visitor log.")

    doc.add_heading("6.6 SOS Panic", level=2)
    doc.add_paragraph(
        "Press Distress SOS to trigger an emergency alert. The screen turns red with a full "
        "SOS PANIC TRIGGERED overlay. The Command Centre receives an immediate alert with sound. "
        "Press Cancel Distress Alarm to clear the alert."
    )

    doc.add_page_break()

    # --- 7 Workflows ---
    doc.add_heading("7. End-to-End Workflows", level=1)

    doc.add_heading("7.1 Daily Patrol Workflow", level=2)
    for i, step in enumerate([
        "Supervisor opens Command Centre and confirms checkpoints are configured.",
        "Guard opens mobile app, verifies online status, and sets their name.",
        "Guard walks the patrol route and taps each NFC checkpoint.",
        "Command Centre map and compliance rate update in real time.",
        "Supervisor reviews OB entries for any missed or overdue checkpoints.",
    ], 1):
        add_numbered(doc, step)
    doc.add_paragraph()

    doc.add_heading("7.2 Emergency SOS Workflow", level=2)
    for i, step in enumerate([
        "Guard presses Distress SOS on mobile.",
        "Red overlay confirms the alert was sent.",
        "Command Centre displays Critical Distress Signal banner with guard name and time.",
        "Siren sounds on supervisor workstation (if enabled).",
        "Supervisor dispatches response and updates OB incident status.",
        "Guard or supervisor dismisses the alarm when resolved.",
    ], 1):
        add_numbered(doc, step)
    doc.add_paragraph()

    doc.add_heading("7.3 Visitor Check-In Workflow", level=2)
    for i, step in enumerate([
        "Visitor arrives at the gate or reception.",
        "Guard scans QR badge or manually enters visitor details on Access tab.",
        "Guard submits the registration.",
        "Supervisor sees the visitor in Active Guests count and visitor log on web.",
        "On departure, supervisor clicks Check Out on the web dashboard.",
    ], 1):
        add_numbered(doc, step)

    doc.add_page_break()

    # --- 8 Offline ---
    doc.add_heading("8. Offline Mode & Synchronisation", level=1)
    doc.add_paragraph(
        "Guards often work in areas with poor connectivity. The mobile app supports full offline operation."
    )
    doc.add_heading("8.1 How Offline Mode Works", level=2)
    for item in [
        "Tap the Wi-Fi icon to switch to offline mode manually, or the app auto-switches if the server is unreachable.",
        "All patrol taps, incidents, visitor check-ins, and checklists are saved locally on the device.",
        "A sync badge shows the number of queued operations.",
        "When connection restores, tap the sync badge or toggle back to online — all queued items upload automatically.",
    ]:
        add_bullet(doc, item)

    # --- 9 Multi-tenant ---
    doc.add_heading("9. Multi-Tenant Management", level=1)
    doc.add_paragraph(
        "Titan Protection supports multiple security providers or client sites as separate tenants. "
        "Each tenant has its own checkpoints, checklists, visitors, and OB entries."
    )
    add_table(doc,
        ["Tenant", "Description"],
        [
            ["Titan Protection", "Default enterprise tenant"],
            ["Alpha Guard Corp", "Secondary security provider"],
            ["Omega Watchmen", "Additional client site"],
        ],
    )
    doc.add_paragraph(
        "On mobile, select the On-Duty Agency Profile in Server Network Config. "
        "On web, use the Active Security Tenant dropdown in the sidebar."
    )

    doc.add_page_break()

    # --- 10 Troubleshooting ---
    doc.add_heading("10. Troubleshooting & FAQ", level=1)
    add_table(doc,
        ["Issue", "Solution"],
        [
            ["Web dashboard shows loading forever", "Ensure the backend server is running. Check browser console for errors."],
            ["Mobile cannot connect to server", "Verify Server URL. Use device IP address, not localhost, on physical phones."],
            ["Data not syncing between mobile and web", "Confirm both apps point to the same backend URL and tenant."],
            ["SOS alert has no sound", "Enable Siren Alerts in the web sidebar. Allow browser audio permissions."],
            ["Offline items not syncing", "Tap the sync badge or toggle Wi-Fi icon to force reconnect."],
            ["Demo Mode badge on web", "Normal in development. Set FORCE_SUPABASE=1 for live database."],
        ],
    )

    doc.add_page_break()

    # --- Appendix Video Script ---
    doc.add_heading("Appendix A: Video Walkthrough Script", level=1)
    doc.add_paragraph(
        "Use this script when recording the system walkthrough video. Open the HTML "
        "presentation (Titan-Protection-System-Walkthrough.html) alongside the live apps."
    )

    script_sections = [
        ("Opening (0:00–0:30)",
         "Welcome to Titan Protection — Built to Protect. This platform gives security "
         "supervisors a real-time Command Centre and equips guards with a powerful mobile app "
         "for patrols, emergencies, and access control."),
        ("Web — Command Centre (0:30–2:00)",
         "The web dashboard is your operations hub. At a glance you see compliance rate, "
         "checkpoint scans, active visitors, and alerts. The live map tracks guard movement "
         "and checkpoint status. The Occurrence Book logs every event — patrol taps, incidents, "
         "and SOS alerts. When a guard triggers SOS, this red banner appears instantly with sound."),
        ("Web — Visitors & Admin (2:00–2:45)",
         "Register visitors from the Access Desk or review the live log. Switch tenants from "
         "the sidebar. In Master Admin, build custom checklist templates and onboard new "
         "security providers."),
        ("Mobile — Patrol & SOS (2:45–4:00)",
         "On the guard mobile app, the Patrol tab shows a radar map and checkpoint list. "
         "Tap NFC Tap at each checkpoint — the Command Centre updates within seconds. "
         "The Distress SOS button sends an immediate emergency alert to supervisors."),
        ("Mobile — Incidents, Checklists, Access (4:00–5:00)",
         "Report incidents with photos and voice memos. Complete compliance checklists with "
         "progress tracking. Check in visitors via QR scan or manual entry."),
        ("Offline & Closing (5:00–5:30)",
         "The app works offline — all actions queue locally and sync when connection returns. "
         "Titan Protection: Built to Protect. Thank you for watching."),
    ]
    for heading, text in script_sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(text)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Titan Protection  |  Built to Protect  |  Arch Luviah Technologies © 2026")
    run.font.size = Pt(9)
    run.font.color.rgb = GREEN

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
