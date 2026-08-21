#!/usr/bin/env python3
"""Generate Titan Protection System Manual (.docx) — whole ecosystem, branded."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUT = DOCS / "Titan-Protection-System-Manual.docx"
OLD_OUTPUT = DOCS / "Titan-Protection-Web-User-Manual.docx"

WORDMARK = ROOT / "web" / "public" / "emblem-wordmark.png"
EMBLEM = ROOT / "web" / "public" / "emblem-light.jpg"
APP_ICON = ROOT / "web" / "public" / "icons" / "icon-512.png"

# Titan Protection brand palette (from web/app/globals.css)
BRAND = {
    "primary": RGBColor(0x1B, 0x43, 0x32),
    "primary_hover": RGBColor(0x2D, 0x6A, 0x4F),
    "primary_light_hex": "D8F3DC",
    "primary_hex": "1B4332",
    "success": RGBColor(0x40, 0x91, 0x6C),
    "success_hex": "40916C",
    "warning_hex": "F59E0B",
    "danger_hex": "EF4444",
    "text": RGBColor(0x0F, 0x1F, 0x17),
    "muted": RGBColor(0x3D, 0x5A, 0x48),
    "dimmed": RGBColor(0x7A, 0x94, 0x85),
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "surface_hex": "F4FAF6",
}


def hex_fill(cell, fill_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BRAND["text"]

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = BRAND["primary"]
        h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(6)


def setup_header_footer(doc: Document, logo: Path):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        # Header (pages 2+)
        hdr = section.header
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if logo.exists():
            hp.add_run().add_picture(str(logo), width=Inches(1.45))
        tr = hp.add_run("   Titan Protection — System Manual")
        tr.font.size = Pt(9)
        tr.font.color.rgb = BRAND["muted"]
        tr.italic = True

        # Footer
        ftr = section.footer
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(
            f"Titan Protection Operations Hub  ·  titanprotection.org  ·  "
            f"Arch Luviah Technologies © {date.today().year}  ·  Confidential"
        )
        r.font.size = Pt(8)
        r.font.color.rgb = BRAND["dimmed"]


def add_cover_page(doc: Document):
    # Brand banner
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.rows[0].cells[0]
    hex_fill(cell, BRAND["primary_hex"])
    cell.height = Cm(1.2)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TITAN PROTECTION")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND["white"]

    doc.add_paragraph()

    # Wordmark logo
    if WORDMARK.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(str(WORDMARK), width=Inches(3.2))
    elif APP_ICON.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(str(APP_ICON), width=Inches(1.6))

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("System Manual")
    r.bold = True
    r.font.size = Pt(32)
    set_run_color(r, BRAND["primary"])

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Complete Operations Platform Guide\n"
        "Web Command Portal  ·  Supervisor Portal  ·  Mobile Field Clients"
    )
    r.font.size = Pt(13)
    set_run_color(r, BRAND["muted"])

    doc.add_paragraph()

    # Accent shield icon
    if APP_ICON.exists():
        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(str(APP_ICON), width=Inches(0.85))

    doc.add_paragraph()

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Document", "Titan Protection System Manual"),
        ("Version", "1.0"),
        ("Date", date.today().strftime("%d %B %Y")),
        ("Live system", "https://titanprotection.org"),
    ]
    for i, (label, val) in enumerate(rows):
        hex_fill(meta.rows[i].cells[0], BRAND["primary_light_hex"])
        c0 = meta.rows[i].cells[0].paragraphs[0]
        c0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = c0.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        set_run_color(lr, BRAND["primary"])
        c1 = meta.rows[i].cells[1].paragraphs[0]
        vr = c1.add_run(val)
        vr.font.size = Pt(10)
        set_run_color(vr, BRAND["text"])

    doc.add_paragraph()
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tagline.add_run("Built to Protect")
    r.italic = True
    r.font.size = Pt(12)
    set_run_color(r, BRAND["success"])

    dev = doc.add_paragraph()
    dev.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dev.add_run("Developed by Arch Luviah Technologies")
    r.font.size = Pt(10)
    set_run_color(r, BRAND["dimmed"])

    doc.add_page_break()


def add_branded_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hex_fill(hdr[i], BRAND["primary_hex"])
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = BRAND["white"]
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            if ri % 2 == 1:
                hex_fill(cell, BRAND["surface_hex"])
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.color.rgb = BRAND["text"]


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.color.rgb = BRAND["text"]


def add_callout(doc, title, body, kind="note"):
    fill = BRAND["primary_light_hex"] if kind == "note" else "FEE2E2"
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    hex_fill(cell, fill)
    p = cell.paragraphs[0]
    tr = p.add_run(f"{title}  ")
    tr.bold = True
    set_run_color(tr, BRAND["primary"] if kind == "note" else RGBColor(0xB9, 0x1C, 0x1C))
    p.add_run(body)
    doc.add_paragraph()


def add_part_banner(doc, part_num, title):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    hex_fill(cell, BRAND["success_hex"])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"PART {part_num}  ·  {title.upper()}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND["white"]
    doc.add_paragraph()


def add_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    parts = [
        ("I", "System Introduction", [
            "1. About Titan Protection",
            "2. System Architecture & Components",
            "3. User Roles & Access Model",
            "4. Data Flow & Synchronisation",
        ]),
        ("II", "Access & Web Portals", [
            "5. Signing In",
            "6. Master Admin Web Portal",
            "7. Supervisor Web Portal",
        ]),
        ("III", "Mobile Field Clients", [
            "8. Titan Monitor — Field Operations Client",
            "9. Titan Supervisor — Mobile Client",
        ]),
        ("IV", "Core Operational Modules", [
            "10. Territory & Supervisor Administration",
            "11. Premises & Site Registration",
            "12. Personnel & Shift Management",
            "13. Command Centre & Real-Time Monitoring",
            "14. GIS Operations Map",
            "15. Data Explorer & Records Management",
        ]),
        ("V", "Platform Services", [
            "16. Communications, PIN Delivery & Alerts",
            "17. System Settings & Master Administration",
            "18. Mobile App Distribution & Updates",
        ]),
        ("VI", "Support", [
            "19. Troubleshooting & FAQ",
            "Appendix A — Quick Reference",
            "Appendix B — Glossary",
        ]),
    ]
    for part, name, items in parts:
        p = doc.add_paragraph()
        r = p.add_run(f"Part {part} — {name}")
        r.bold = True
        set_run_color(r, BRAND["primary"])
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


def build_content(doc):
    # === PART I ===
    add_part_banner(doc, "I", "System Introduction")

    doc.add_heading("1. About Titan Protection", level=1)
    doc.add_paragraph(
        "Titan Protection is an integrated security operations platform designed for private "
        "security companies, corporate protection teams, and institutional guard services. "
        "The system unifies command-and-control, field operations, geographic site management, "
        "personnel administration, and real-time incident response in a single live ecosystem."
    )
    doc.add_paragraph(
        "This manual documents the complete Titan Protection system — not a single application "
        "or user role, but the full platform: web command portals, supervisor workspaces, "
        "Android field clients, shared database, and operational workflows that connect them."
    )
    doc.add_heading("1.1 System Purpose", level=2)
    add_bullets(doc, [
        "Register and geolocate protected premises and patrol points.",
        "Organise operations by territory with dedicated area supervisors.",
        "Deploy and monitor field personnel with GPS-verified attendance.",
        "Track patrol activity, incidents, and compliance in real time.",
        "Respond to distress (SOS) signals and operational alerts immediately.",
        "Maintain audit trails through the Occurrence Book and Data Explorer.",
        "Distribute mobile clients and keep all endpoints synchronised.",
    ])
    doc.add_heading("1.2 Who Uses This Manual", level=2)
    add_branded_table(doc, ["Audience", "Interest"], [
        ["Master Administrators", "Full platform setup, configuration, and governance"],
        ["Operations Managers", "Daily monitoring via Command Centre and GIS map"],
        ["Area Supervisors", "Territory-scoped web and mobile supervisor tools"],
        ["IT / Deployment Teams", "APK distribution, login configuration, database sync"],
        ["Training & Compliance Officers", "Workflows, checklists, audit records"],
    ])

    doc.add_heading("2. System Architecture & Components", level=1)
    doc.add_paragraph(
        "Titan Protection consists of four client applications sharing one API and database. "
        "All components read from and write to the same operational datastore when connected "
        "to the live server (Supabase)."
    )
    add_branded_table(doc, ["Component", "Platform", "Primary function"], [
        ["Master Admin Web Portal", "Browser — /", "Full system administration and monitoring"],
        ["Supervisor Web Portal", "Browser — /supervisor", "Territory-scoped operations dashboard"],
        ["Titan Monitor", "Android — com.titan.monitor", "Field operations client for on-site personnel"],
        ["Titan Supervisor", "Android — com.titan.supervisor", "Field client for supervisors — GPS site capture, team view"],
    ])
    doc.add_heading("2.1 Shared Infrastructure", level=2)
    add_bullets(doc, [
        "API layer — /api/state handles reads and operational mutations for all clients.",
        "Authentication — separate admin (email/password) and supervisor (6-digit PIN) sessions.",
        "Database — Supabase PostgreSQL; demo mode uses in-memory store when server unreachable.",
        "Real-time polling — web dashboards refresh every 10 seconds; mobile clients poll on comparable intervals.",
        "Production URL — https://titanprotection.org (Vercel deployment, web root: web/).",
    ])
    doc.add_heading("2.2 Recommended Deployment Sequence", level=2)
    add_numbered(doc, [
        "Configure Master Admin credentials and connect to live database.",
        "Create territories and register supervisors; deliver PINs.",
        "Register premises with on-site GPS capture (±5 m accuracy).",
        "Add patrol places at each premises.",
        "Register field personnel; assign one supervisor and premises per person.",
        "Schedule shifts and distribute Titan Monitor APK.",
        "Supervisors install Titan Supervisor APK for field GPS capture.",
        "Monitor operations in Command Centre and GIS map.",
    ])

    doc.add_heading("3. User Roles & Access Model", level=1)
    add_branded_table(doc, ["Role", "Portal / App", "Scope"], [
        ["Master Admin", "Web — /", "Entire tenant — all territories, settings, data"],
        ["Area Supervisor", "Web — /supervisor + Titan Supervisor app", "Assigned territories only"],
        ["Field Officer", "Titan Monitor app", "Assigned premises and shifts"],
    ])
    doc.add_heading("3.1 Supervisor Territory Scoping", level=2)
    doc.add_paragraph(
        "Supervisors (web and mobile) see only data linked to their assigned territory IDs: "
        "premises in those territories, places under those premises, guards assigned to them "
        "or their territories, related shifts, attendance, alerts, and occurrence book entries."
    )
    add_callout(doc, "Access rule:", "Supervisors cannot create territories, manage other supervisors, "
                 "reassign guards to another supervisor, or access Master Admin tools.")

    doc.add_heading("4. Data Flow & Synchronisation", level=1)
    doc.add_paragraph("Typical operational data flow across the system:")
    add_numbered(doc, [
        "Admin or supervisor registers a premises with GPS → stored in database → appears on GIS map (green marker).",
        "Patrol places added at premises → small red markers on map; checkpoint records sync for mobile patrol.",
        "Field officer clocks in via Titan Monitor → attendance record → Command Centre and map update.",
        "Incident or SOS logged on mobile → Occurrence Book entry → alert banner on web dashboards.",
        "Supervisor captures site GPS on Titan Supervisor app → premises/places update via API.",
        "All web dashboards poll /api/state; changes propagate within seconds on live connection.",
    ])
    add_callout(doc, "Connection indicator:", "Server Connected / Live = production database. "
                 "Demo Mode = in-memory only; data may not persist.")

    # === PART II ===
    add_part_banner(doc, "II", "Access & Web Portals")

    doc.add_heading("5. Signing In", level=1)
    doc.add_heading("5.1 Login Chooser — /login", level=2)
    doc.add_paragraph("Welcome to Titan Protection — choose your role:")
    add_branded_table(doc, ["Role card", "Description", "Action"], [
        ["Master Admin", "Full access — supervisors, territories, personnel, premises, settings", "Admin sign in →"],
        ["Supervisor", "Area-scoped — command centre, personnel, premises in assigned territories", "Supervisor sign in →"],
    ])
    doc.add_paragraph("Footer: Download Titan Monitor & Supervisor APKs → /downloads")

    doc.add_heading("5.2 Master Admin — /admin/login", level=2)
    add_numbered(doc, [
        "Enter email and password configured by your system administrator.",
        "Click Sign In → redirected to / (Master Admin Dashboard).",
        "Sign Out from the page header when finished.",
    ])

    doc.add_heading("5.3 Supervisor — /supervisor/login", level=2)
    add_numbered(doc, [
        "Enter 6-digit PIN on the numeric keypad (same PIN as Titan Supervisor mobile app).",
        "Redirected to /supervisor (Supervisor Dashboard).",
        "Forgotten PIN: Master Admin resets from Supervisor & Territories → Reset PIN.",
    ])

    doc.add_heading("6. Master Admin Web Portal", level=1)
    doc.add_paragraph("Full operations hub. Sidebar navigation:")
    add_branded_table(doc, ["Module", "Function"], [
        ["Supervisor & Territories", "Territories, supervisor accounts, WhatsApp setup"],
        ["Guard Management", "Personnel profiles, shifts, attendance, alerts, performance"],
        ["Register Premises", "Protected sites, patrol places, GPS capture"],
        ["GIS Operations Map", "Live geographic operations view"],
        ["Command Centre", "Real-time stats, map, Occurrence Book, visitors"],
        ["Data Explorer", "Database collections — search, export, edit"],
        ["Master Admin", "Sync, checklists, audits, data management"],
        ["Mobile App Downloads", "APK distribution page"],
    ])
    doc.add_heading("6.1 System Settings (sidebar)", level=2)
    add_branded_table(doc, ["Setting group", "Options"], [
        ["Alerts & Sirens", "Siren on critical alerts (SOS, critical guard alerts)"],
        ["GPS & Geofencing", "Geofence radius 5–8 m (default 6 m)"],
        ["Patrol Monitoring", "No-movement threshold 15–90 min (default 45 min)"],
        ["License Compliance", "Expiry warning window 14–90 days (default 60 days)"],
    ])
    doc.add_heading("6.2 Connection & Refresh", level=2)
    doc.add_paragraph(
        "Header badge shows Server Connected or Demo Mode Active. Dashboard polls every 10 seconds. "
        "Critical alerts may trigger audible siren when enabled."
    )

    doc.add_heading("7. Supervisor Web Portal", level=1)
    doc.add_paragraph("Territory-scoped mirror of key admin modules:")
    add_branded_table(doc, ["Module", "Supervisor scope"], [
        ["Register Premises", "Premises in assigned territories"],
        ["Guard Management", "Scoped personnel only"],
        ["GIS Operations Map", "Scoped sites, places, events"],
        ["Command Centre", "Scoped live monitoring"],
        ["Data Explorer", "Scoped records; limited edit tables"],
    ])
    add_callout(doc, "Not available to supervisors:", "Supervisor & Territories admin, Master Admin tab, "
                 "System Settings, Mobile Downloads link, Reset Patrol, cross-territory data.")

    # === PART III ===
    add_part_banner(doc, "III", "Mobile Field Clients")

    doc.add_heading("8. Titan Monitor — Field Operations Client", level=1)
    doc.add_paragraph(
        "Package: com.titan.monitor · Android APK from /downloads\n"
        "Purpose: Primary field client for on-site security personnel."
    )
    doc.add_heading("8.1 Sign-In & Profile", level=2)
    add_bullets(doc, [
        "6-digit PIN login (delivered by email on registration).",
        "First-login PIN change flow when required.",
        "Profile photo upload; assigned premises resolved from guard profile and active shift.",
        "Over-the-air update prompt when newer APK published.",
    ])
    doc.add_heading("8.2 Main Modules (bottom navigation)", level=2)
    add_branded_table(doc, ["Tab", "Function"], [
        ["Patrol", "Clock-in/out with GPS geofence verification; patrol point progress; movement tracking"],
        ["Incidents", "Log incidents with photo and voice memo attachments; syncs to Occurrence Book"],
        ["Checklists", "Complete published audit/checklist templates from web admin"],
        ["Access", "Visitor QR scanning and access desk functions"],
        ["Profile", "Guard details, documents, performance summary, sign out"],
    ])
    doc.add_heading("8.3 Key Field Operations", level=2)
    add_bullets(doc, [
        "GPS clock-in — must be within premises geofence radius (system setting, default 6 m).",
        "Patrol progress — tracks scanned/completed patrol points for assigned site.",
        "SOS distress — triggers CRITICAL DISTRESS SIGNAL on all web Command Centres.",
        "No-movement monitoring — server-side alerts if guard stationary beyond threshold.",
        "Geofence exit alerts — triggered when on-duty guard leaves site boundary.",
    ])

    doc.add_heading("9. Titan Supervisor — Mobile Client", level=1)
    doc.add_paragraph(
        "Package: com.titan.supervisor · Android APK from /downloads\n"
        "Purpose: Field tool for area supervisors — complements the supervisor web portal."
    )
    doc.add_heading("9.1 Sign-In", level=2)
    doc.add_paragraph(
        "Same 6-digit PIN as supervisor web portal. Session scoped to supervisor ID and assigned territories."
    )
    doc.add_heading("9.2 Main Modules", level=2)
    add_branded_table(doc, ["Tab", "Function"], [
        ["Home", "Dashboard summary — territories, premises, guards, alerts"],
        ["Sites", "Register premises and patrol places with high-accuracy GPS capture (±5 m target)"],
        ["Team", "View and manage scoped guard team"],
        ["Profile", "Supervisor photo, PIN change, app update, sign out"],
    ])
    doc.add_heading("9.3 GPS Capture", level=2)
    doc.add_paragraph(
        "Sites panel uses extended GPS warm-up and cluster-median filtering for accurate coordinates. "
        "Save is optimised for fast completion after GPS lock. Captured coordinates sync to web GIS map."
    )

    # === PART IV ===
    add_part_banner(doc, "IV", "Core Operational Modules")

    doc.add_heading("10. Territory & Supervisor Administration", level=1)
    doc.add_paragraph("Master Admin only — Supervisor & Territories module.")
    doc.add_heading("10.1 Territories", level=2)
    add_numbered(doc, [
        "Add Territory → name, city, description, suburbs list.",
        "Territories scope all premises, supervisors, and filtered data views.",
        "Delete only after reassigning linked premises and personnel.",
    ])
    doc.add_heading("10.2 Supervisors", level=2)
    add_numbered(doc, [
        "Add Supervisor → name, role, phone, email, assigned territories.",
        "System generates 6-digit PIN; delivered via email + WhatsApp manual send.",
        "Reset PIN, upload photo, edit status (Active / Suspended / Off Duty).",
    ])
    doc.add_heading("10.3 WhatsApp Integration", level=2)
    doc.add_paragraph(
        "Manual WhatsApp model — system opens WhatsApp with pre-filled message; user taps Send. "
        "Used for PIN delivery backup, shift notifications, and supervisor-to-guard messaging."
    )

    doc.add_heading("11. Premises & Site Registration", level=1)
    doc.add_paragraph(
        "Available on Master Admin and Supervisor web portals (supervisor: scoped territories). "
        "Also via Titan Supervisor mobile Sites tab."
    )
    doc.add_heading("11.1 Register Premises", level=2)
    add_numbered(doc, [
        "Premises name, owner/client, address, territory, city, suburb.",
        "Capture GPS on site — target ±5 m accuracy.",
        "Share Premises ID with field personnel for Titan Monitor site linking.",
    ])
    doc.add_heading("11.2 Patrol Places", level=2)
    add_numbered(doc, [
        "Add Place under selected premises — name, type, schedule, description.",
        "Capture place GPS at physical location.",
        "Map display: small red P markers. NFC tagging deferred for future release.",
    ])
    doc.add_heading("11.3 Map Markers (current)", level=2)
    add_branded_table(doc, ["Marker", "Colour / size", "Meaning"], [
        ["Premises", "Green — site initials", "Registered protected site centre"],
        ["Patrol place", "Red — small P pin", "Important location inside a premises"],
    ])

    doc.add_heading("12. Personnel & Shift Management", level=1)
    doc.add_paragraph("Guard Management module — personnel administration (web).")
    add_branded_table(doc, ["Sub-tab", "Purpose"], [
        ["Guard Profiles", "Register personnel, documents, bulk supervisor assignment"],
        ["Shift Roster", "Schedule shifts; WhatsApp notification on save"],
        ["Live Attendance", "On-duty status, geofence/movement badges, history"],
        ["Supervisor Alerts", "Movement, geofence, license alerts — dismiss when resolved"],
        ["Performance", "Composite scorecards"],
        ["Shift Swaps", "Approve/reject mobile swap requests"],
        ["Messaging", "Email and WhatsApp configuration and send"],
    ])
    doc.add_heading("12.1 Personnel Registration Rules", level=2)
    add_bullets(doc, [
        "Each field officer assigned to exactly one supervisor (supervisorId).",
        "Supervisors registering new personnel must assign them to themselves.",
        "Bulk tools: Assign all to supervisor; Auto-match by territory.",
        "PIN delivered via email for Titan Monitor login.",
    ])

    doc.add_heading("13. Command Centre & Real-Time Monitoring", level=1)
    doc.add_paragraph("Shared by Master Admin and Supervisor portals (supervisor: scoped data).")
    add_branded_table(doc, ["Panel", "Content"], [
        ["SOS banner", "CRITICAL DISTRESS SIGNAL — Dismiss Alarm"],
        ["Statistics", "Compliance rate, taps scanned, active guests, guards on duty"],
        ["Live map", "Compact GIS view of sites and patrol points"],
        ["Occurrence Book", "Chronological field events — status workflow, photo/voice attachments"],
        ["Visitor registry", "Check-in form and live visitor log table"],
    ])

    doc.add_heading("14. GIS Operations Map", level=1)
    doc.add_paragraph("Full-screen map (Leaflet) with sidebar in dedicated tab; compact embed in Command Centre.")
    doc.add_heading("14.1 Basemaps", level=2)
    add_bullets(doc, [
        "Google Satellite Hybrid (default), Google Streets, Google Terrain, Dark Matter GIS, Esri World Imagery.",
    ])
    doc.add_heading("14.2 View Presets", level=2)
    add_branded_table(doc, ["Preset", "Use case"], [
        ["Operations", "Live personnel, geofences, trails, alerts"],
        ["Patrol", "Premises and patrol places only (simplified default)"],
        ["Emergency", "Alerts, heatmap, activity focus"],
        ["Overview", "Sites and places across territories"],
    ])
    doc.add_heading("14.3 Layer Toggles", level=2)
    add_bullets(doc, [
        "Protected premises, patrol places, geofences, GPS accuracy rings, patrol routes.",
        "Live personnel, movement trails, shift roster, alerts, activity, heatmap, territory zones.",
        "NFC checkpoint layer hidden until NFC patrol map integration is enabled.",
    ])
    doc.add_heading("14.4 Map Tools", level=2)
    add_bullets(doc, [
        "Search, zoom, fit all sites, measure distance, copy GPS from popups, fly-to on SOS.",
    ])

    doc.add_heading("15. Data Explorer & Records Management", level=1)
    doc.add_paragraph(
        "Browse all system collections — Core, Geography, People, Operations, Communications, "
        "Compliance, Emergency categories."
    )
    add_bullets(doc, [
        "Search, paginate (30 rows), copy JSON, export CSV/JSON.",
        "Inline edit (admin: all editable tables; supervisor: premises, places, guards, shifts, alerts, OB).",
        "Sensitive fields redacted — PINs, passwords, raw GPS trails, attachments.",
    ])

    # === PART V ===
    add_part_banner(doc, "V", "Platform Services")

    doc.add_heading("16. Communications, PIN Delivery & Alerts", level=1)
    add_branded_table(doc, ["Event", "Email", "WhatsApp"], [
        ["New field officer registered", "Titan Monitor PIN", "Manual send modal"],
        ["New supervisor registered", "Supervisor PIN", "Manual send modal"],
        ["PIN reset", "New PIN emailed", "WhatsApp opens"],
        ["Shift scheduled", "—", "Notification message"],
        ["SOS triggered", "—", "Command Centre banner + optional siren"],
    ])
    doc.add_heading("16.1 Alert Types", level=2)
    add_bullets(doc, [
        "SOS panic — highest priority; red banner all dashboards.",
        "No-movement — field officer stationary beyond threshold.",
        "Geofence exit — on-duty officer left site boundary.",
        "License expiry — compliance warning window.",
        "Shift swap request — pending supervisor action.",
    ])

    doc.add_heading("17. System Settings & Master Administration", level=1)
    doc.add_heading("17.1 Master Admin Tab", level=2)
    add_bullets(doc, [
        "Reload from database — sync from Supabase.",
        "Clear all data — destructive demo reset (production caution).",
        "Checklist Template Builder — publish audit forms for Titan Monitor.",
        "Completed Audits — view submitted checklist responses.",
        "System Configuration summary — read-only settings display.",
    ])
    doc.add_heading("17.2 Security", level=2)
    add_bullets(doc, [
        "Role-isolated sessions — admin vs supervisor cookies.",
        "PIN redaction in Data Explorer and supervisor API responses.",
        "Supervisor mutation guard — server validates territory scope on every action.",
        "HTTPS enforced in production.",
    ])

    doc.add_heading("18. Mobile App Distribution & Updates", level=1)
    doc.add_paragraph("Public page: /downloads")
    add_branded_table(doc, ["App", "Package", "Audience"], [
        ["Titan Monitor", "com.titan.monitor", "Field operations personnel"],
        ["Titan Supervisor", "com.titan.supervisor", "Area supervisors"],
    ])
    doc.add_heading("18.1 Installation", level=2)
    add_numbered(doc, [
        "Download APK from /downloads.",
        "Allow installation from unknown sources (Android).",
        "Sign in with PIN.",
        "Use in-app Update app when newer build published — OTA without manual reinstall.",
    ])

    # === PART VI ===
    add_part_banner(doc, "VI", "Support")

    doc.add_heading("19. Troubleshooting & FAQ", level=1)
    add_branded_table(doc, ["Issue", "Resolution"], [
        ["Demo Mode showing", "Verify Supabase connection; Reload from database"],
        ["Wrong site on Titan Monitor", "Verify personnel premises assignment; logout/login after app update"],
        ["0/0 patrol points", "Ensure places exist under premises; checkpoints sync from places"],
        ["GPS capture fails", "Enable location; stand outdoors; wait for ±5 m lock"],
        ["Blue NFC markers on map", "Disable checkpoints layer; NFC map integration deferred"],
        ["Supervisor PIN failed", "Master Admin → Reset PIN"],
        ["Data not updating", "Wait 10 s auto-refresh or use Refresh in Data Explorer"],
    ])
    doc.add_heading("19.1 FAQ", level=2)
    add_branded_table(doc, ["Question", "Answer"], [
        ["Is this a guard-only manual?", "No — this documents the entire Titan Protection system for all roles."],
        ["Can one officer have two supervisors?", "No — exactly one supervisor per field officer."],
        ["Do supervisors see all company data?", "No — only assigned territories."],
        ["How many map marker types now?", "Two — green premises, red patrol places."],
        ["Where are APKs?", "/downloads or login page footer link."],
    ])

    doc.add_heading("Appendix A — Quick Reference", level=1)
    add_branded_table(doc, ["Resource", "URL / path"], [
        ["Login chooser", "/login"],
        ["Admin dashboard", "/"],
        ["Supervisor dashboard", "/supervisor"],
        ["Mobile downloads", "/downloads"],
        ["Production", "https://titanprotection.org"],
    ])
    doc.add_heading("System Onboarding Checklist", level=2)
    add_numbered(doc, [
        "Connect live database; confirm Server Connected.",
        "Create territories and supervisors.",
        "Register premises with GPS.",
        "Add patrol places.",
        "Register personnel; assign supervisor and premises.",
        "Schedule shifts.",
        "Deploy Titan Monitor and Titan Supervisor APKs.",
        "Verify Command Centre and GIS map.",
    ])

    doc.add_heading("Appendix B — Glossary", level=1)
    add_branded_table(doc, ["Term", "Definition"], [
        ["Titan Protection System", "The complete platform — web portals, mobile apps, API, and database"],
        ["Premises", "A registered protected site with GPS coordinates"],
        ["Patrol place", "Important location inside a premises (red map marker)"],
        ["Territory", "Geographic operating area (city + suburbs)"],
        ["Command Centre", "Real-time web monitoring module"],
        ["Occurrence Book (OB)", "System-wide chronological event log"],
        ["Titan Monitor", "Android field operations client"],
        ["Titan Supervisor", "Android supervisor field client"],
        ["Geofence", "GPS radius for clock-in verification (5–8 m)"],
        ["Compliance rate", "Patrol checkpoint scan completion percentage"],
    ])

    doc.add_paragraph()
    closing = doc.add_table(rows=1, cols=1)
    cell = closing.rows[0].cells[0]
    hex_fill(cell, BRAND["primary_hex"])
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run("\nBuilt to Protect\n")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = BRAND["white"]
    r2 = cp.add_run(f"Titan Protection System Manual v1.0  ·  {date.today().strftime('%B %Y')}\n")
    r2.font.size = Pt(10)
    r2.font.color.rgb = BRAND["white"]
    r3 = cp.add_run("Developed by Arch Luviah Technologies")
    r3.font.size = Pt(9)
    r3.font.color.rgb = BRAND["white"]


def build_manual():
    doc = Document()
    setup_styles(doc)
    logo = WORDMARK if WORDMARK.exists() else APP_ICON
    setup_header_footer(doc, logo)
    add_cover_page(doc)
    add_toc(doc)
    build_content(doc)

    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    if OLD_OUTPUT.exists() and OLD_OUTPUT != OUTPUT:
        OLD_OUTPUT.unlink()
    print(f"Written: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_manual()
